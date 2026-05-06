import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_technical_doc(filename="Support_Ticket_Router_Technical_Document.docx"):
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    title = doc.add_heading('Technical Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('Support Ticket Router', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')

    # Introduction
    doc.add_heading('1. Introduction', level=2)
    doc.add_paragraph(
        "Support Ticket Router is an intelligent, production-grade engine designed to automatically "
        "classify, prioritize, and route customer support tickets using rule-based logic and a full "
        "microservices architecture. This document provides a comprehensive technical overview of the "
        "codebase, architecture, design patterns, and deployment strategies."
    )

    # Architecture Overview
    doc.add_heading('2. Architecture Overview & Tech Stack', level=2)
    p = doc.add_paragraph("The system uses a decoupled microservices architecture with the following core components:")
    components = [
        "API Framework: FastAPI + Uvicorn for high-performance async REST APIs.",
        "Async Processing: Celery for background tasks, using Redis as a message broker and result backend.",
        "Database: Dual-backend system. Defaults to PostgreSQL for production environments and falls back to SQLite for local development or lightweight cloud deployments.",
        "Web Interface: Streamlit providing both an end-user ticket submission interface and an Admin Panel for live rule management.",
        "Rate Limiting: SlowAPI is integrated into FastAPI to prevent API abuse.",
        "Containerization: Docker and Docker Compose orchestrate 5 distinct services (PostgreSQL, Redis, API, Worker, CLI)."
    ]
    for comp in components:
        doc.add_paragraph(comp, style='List Bullet')

    # Core Modules
    doc.add_heading('3. Codebase Analysis (Core Modules)', level=2)

    # 3.1 Data Models
    doc.add_heading('3.1 Data Models (models/ticket.py)', level=3)
    doc.add_paragraph(
        "The project uses Python Dataclasses to enforce strong typing. The main models are Ticket and ProcessedTicket. "
        "Ticket defines the schema for incoming requests (id, subject, message, customer_type, created_at). "
        "ProcessedTicket extends this with output properties: category, priority, assigned_team, and reason."
    )

    # 3.2 Evaluation Engine
    doc.add_heading('3.2 Classification Logic (engine/evaluator.py)', level=3)
    doc.add_paragraph(
        "TicketEvaluator handles the core logic. It reads category rules and urgency keywords. "
        "It uses word-boundary Regex matching (\\b) to ensure partial word matches do not cause false positives "
        "(e.g., 'non-refundable' will not trigger the 'refund' rule). "
        "Priority is assigned based on keyword urgency, category context, and whether the customer is 'Premium'. "
        "A detailed reason is generated outlining why a specific classification was chosen."
    )

    # 3.3 Router Engine
    doc.add_heading('3.3 Team Routing (engine/router.py)', level=3)
    doc.add_paragraph(
        "TeamRouter implements a dictionary-based mapping to route categorized tickets to appropriate teams "
        "(e.g., billing -> payments-team). It ensures default fallbacks (general-support) if no specific category matches."
    )

    # 3.4 Database Layer
    doc.add_heading('3.4 Database Layer (database/db.py)', level=3)
    doc.add_paragraph(
        "Implements a dual-backend Factory pattern. It dynamically detects whether the DATABASE_URL points to PostgreSQL "
        "or if it should fall back to a local SQLite file. This script is responsible for database initialization, schema seeding "
        "(categories, priority_rules, processed_tickets tables), and executing CRUD operations for the Admin Panel. "
        "Dependency injection allows the live Streamlit dashboard to update rules dynamically without application restarts."
    )

    # 3.5 API Layer
    doc.add_heading('3.5 REST API (api.py)', level=3)
    doc.add_paragraph(
        "The FastAPI app provides asynchronous endpoints for processing tickets. "
        "Endpoints include /api/v1/process-ticket for asynchronous enqueueing via Celery, and /api/v1/task/{task_id} "
        "to poll the task status. Pydantic is used for request validation, ensuring bad data is rejected before processing. "
        "SlowAPI limits requests to 30/min for POST and 120/min for GET."
    )

    # 3.6 Background Worker
    doc.add_heading('3.6 Background Processing (worker.py)', level=3)
    doc.add_paragraph(
        "The Celery worker decouples heavy classification from the main thread. It retrieves tasks from Redis, initializes "
        "the TicketEvaluator and TeamRouter via the database rules, processes the data, saves the outcome to the DB history, "
        "and triggers optional Discord/Slack webhooks via the notifier module."
    )

    # 3.7 UI and Admin Panel
    doc.add_heading('3.7 Streamlit Application (app.py)', level=3)
    doc.add_paragraph(
        "The Streamlit app acts as the primary web client. It offers two main sections: a Dashboard to simulate "
        "ticket creation and view real-time historical results, and an Admin Panel. The Admin Panel is password-protected "
        "and offers direct CRUD interfaces to the database to add, delete, or modify category mappings and keywords."
    )

    # Security & Scalability
    doc.add_heading('4. Security & Scalability', level=2)
    scalability_points = [
        "Environment Variables: Secrets are managed securely via .env, parsed by config/settings.py.",
        "Input Sanitization: Pydantic and safe database parameterization prevent SQL injection.",
        "Async Processing: Celery/Redis prevents long blocking operations on the main FastAPI thread, maximizing throughput.",
        "Rate Limiting: IP-based limits protect API endpoints from DoS and scraping.",
        "Multi-Stage Docker Builds: Ensures lightweight production containers with clean dependencies."
    ]
    for point in scalability_points:
        doc.add_paragraph(point, style='List Bullet')

    # Testing
    doc.add_heading('5. Testing', level=2)
    doc.add_paragraph(
        "The system has 26 automated Pytest tests under the tests/ directory. These encompass unit tests for core "
        "engine components (e.g., regex matching edge cases) and comprehensive End-to-End pipeline tests to ensure data flows "
        "accurately from API input to DB persistence."
    )

    doc.add_paragraph("\nDocument generated automatically based on codebase analysis.")
    doc.save(filename)
    print(f"Document saved to {filename}")

if __name__ == '__main__':
    create_technical_doc()
